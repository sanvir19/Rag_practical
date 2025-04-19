from PyPDF2 import PdfReader
import docx
import fitz
from PIL import Image
import pytesseract
import io
import os
from typing import List, Dict, Any
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from ..config import Config

class FileProcessor:
    TEXT_SPLITTER = RecursiveCharacterTextSplitter(
        chunk_size=Config.CHUNK_SIZE,
        chunk_overlap=Config.CHUNK_OVERLAP,
        separators=["\n\n", "\n", " ", ""]
    )

    @staticmethod
    def process_pdf_with_pypdf(pdf_path: str) -> List[Dict[str, Any]]:
        pdf_reader = PdfReader(pdf_path)
        pages_data = []
        for page_num, page in enumerate(pdf_reader.pages, start=1):
            text = page.extract_text()
            if text and text.strip():
                pages_data.append({
                    "text": text,
                    "metadata": {
                        "source": os.path.basename(pdf_path),
                        "page": page_num,
                        "file_path": pdf_path,
                        "method": "pypdf2"
                    }
                })
        return pages_data

    @staticmethod
    def process_pdf_with_ocr(pdf_path: str) -> List[Dict[str, Any]]:
        doc = fitz.open(pdf_path)
        pages_data = []
        for page_num, page in enumerate(doc, start=1):
            pix = page.get_pixmap()
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            text = pytesseract.image_to_string(img)
            pages_data.append({
                "text": text,
                "metadata": {
                    "source": os.path.basename(pdf_path),
                    "page": page_num,
                    "file_path": pdf_path,
                    "method": "ocr"
                }
            })
        return pages_data

    @staticmethod
    def process_docx(docx_path: str) -> List[Dict[str, Any]]:
        doc = docx.Document(docx_path)
        full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text)
        return [{
            "text": full_text,
            "metadata": {
                "source": os.path.basename(docx_path),
                "page": 1,
                "file_path": docx_path,
                "method": "docx"
            }
        }]

    @staticmethod
    def process_txt(txt_path: str) -> List[Dict[str, Any]]:
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                return [{
                    "text": file.read(),
                    "metadata": {
                        "source": os.path.basename(txt_path),
                        "page": 1,
                        "file_path": txt_path,
                        "method": "txt"
                    }
                }]
        except Exception as e:
            print(f"[ERROR] TXT file reading failed: {e}")
            return []

    @staticmethod
    def process_file(file_path: str) -> List[Document]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        if file_path.lower().endswith('.pdf'):
            try:
                file_data = FileProcessor.process_pdf_with_pypdf(file_path)
                if not file_data:
                    file_data = FileProcessor.process_pdf_with_ocr(file_path)
            except Exception as e:
                print(f"[WARN] Falling back to OCR for {file_path}: {e}")
                file_data = FileProcessor.process_pdf_with_ocr(file_path)
        elif file_path.lower().endswith(('.doc', '.docx')):
            file_data = FileProcessor.process_docx(file_path)
        elif file_path.lower().endswith('.txt'):
            file_data = FileProcessor.process_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path}")

        documents = []
        for item in file_data:
            chunks = FileProcessor.TEXT_SPLITTER.split_text(item["text"])
            for chunk in chunks:
                documents.append(Document(
                    page_content=chunk,
                    metadata=item["metadata"].copy()
                ))

        return documents
